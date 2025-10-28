import streamlit as st
import sqlite3
from werkzeug.security import generate_password_hash, check_password_hash
from transformers import pipeline
from docx import Document
import uuid, io, re
from datetime import datetime

# -----------------------------
# DATABASE SETUP
# -----------------------------
conn = sqlite3.connect('econsult.db', check_same_thread=False)
c = conn.cursor()

c.execute('''CREATE TABLE IF NOT EXISTS users
             (id INTEGER PRIMARY KEY AUTOINCREMENT,
              username TEXT UNIQUE,
              password_hash TEXT,
              mobile TEXT,
              created_at TEXT)''')

c.execute('''CREATE TABLE IF NOT EXISTS comments
             (id INTEGER PRIMARY KEY AUTOINCREMENT,
              user_id INTEGER,
              sector TEXT,
              comment TEXT,
              sentiment TEXT,
              summary TEXT,
              passcode TEXT UNIQUE,
              status TEXT,
              created_at TEXT)''')
conn.commit()

# -----------------------------
# BASIC NLP HELPERS
# -----------------------------
positive_words = ["good", "great", "appreciate", "thank", "satisfied", "resolved", "excellent", "helpful"]
negative_words = ["bad", "delay", "urgent", "problem", "stop", "not working", "poor", "issue"]

_sentiment_pipe, _summarizer = None, None

def load_models():
    global _sentiment_pipe, _summarizer
    try:
        _sentiment_pipe = pipeline("sentiment-analysis")
        _summarizer = pipeline("summarization")
    except Exception:
        _sentiment_pipe, _summarizer = None, None

def predict_sentiment(text, use_model=False):
    if use_model and _sentiment_pipe:
        try:
            result = _sentiment_pipe(text[:512])[0]
            return result['label'].lower(), result['score']
        except Exception:
            pass
    # fallback simple keyword-based
    t = text.lower()
    if any(w in t for w in negative_words): return "negative", 0.9
    if any(w in t for w in positive_words): return "positive", 0.9
    return "neutral", 0.6

def summarize_text(text, use_model=False):
    if use_model and _summarizer:
        try:
            out = _summarizer(text, max_length=40, min_length=5, do_sample=False)[0]
            return out['summary_text']
        except Exception:
            pass
    # fallback: first sentence or short truncation
    s = re.split(r'\.|\n', text)
    return (s[0] or text)[:80]

# -----------------------------
# DB HELPERS
# -----------------------------
def create_user(username, password, mobile):
    try:
        ph = generate_password_hash(password)
        c.execute("INSERT INTO users (username, password_hash, mobile, created_at) VALUES (?,?,?,?)",
                  (username, ph, mobile, datetime.now()))
        conn.commit()
        return True, "User created successfully."
    except Exception as e:
        return False, str(e)

def authenticate_user(username, password):
    c.execute("SELECT id, password_hash FROM users WHERE username=?", (username,))
    row = c.fetchone()
    if not row:
        return False, "User not found"
    uid, ph = row
    if check_password_hash(ph, password):
        return True, uid
    else:
        return False, "Incorrect password"

def add_comment(user_id, sector, comment, sentiment, summary):
    code = str(uuid.uuid4())[:8]
    c.execute("INSERT INTO comments (user_id, sector, comment, sentiment, summary, passcode, status, created_at) VALUES (?,?,?,?,?,?,?,?)",
              (user_id, sector, comment, sentiment, summary, code, "Submitted", datetime.now()))
    conn.commit()
    return code

def get_comment_by_passcode(code):
    c.execute("SELECT sector, comment, sentiment, summary, status FROM comments WHERE passcode=?", (code,))
    row = c.fetchone()
    if not row: return None
    keys = ["sector", "comment", "sentiment", "summary", "status"]
    return dict(zip(keys, row))

def list_comments_for_user(uid):
    c.execute("SELECT sector, comment, sentiment, summary, passcode, status FROM comments WHERE user_id=? ORDER BY id DESC", (uid,))
    rows = c.fetchall()
    cols = ["sector","comment","sentiment","summary","passcode","status"]
    return [dict(zip(cols,r)) for r in rows]

# -----------------------------
# STREAMLIT UI
# -----------------------------
st.set_page_config(page_title="MCA E-Consultation Portal", layout="wide")

st.markdown("""
<style>
.reportview-container { background-color: #f7fbff; }
.stButton>button { background-color: #0b66c3; color: white; border-radius: 8px; }
</style>
""", unsafe_allow_html=True)

st.image("https://upload.wikimedia.org/wikipedia/commons/4/4f/Emblem_of_India.svg", width=60)
st.title("🇮🇳 Ministry of Corporate Affairs — E-Consultation Portal")
st.caption("Smart India Hackathon 2025 | AI-powered Sentiment Analysis Module")

if 'user_id' not in st.session_state:
    st.session_state.user_id = None
    st.session_state.username = None

# Sidebar: Login/Register
st.sidebar.header("🔐 User Access")

if st.session_state.user_id is None:
    act = st.sidebar.radio("Choose", ["Login", "Register", "Guest"])
    if act == "Register":
        u = st.sidebar.text_input("Username")
        p = st.sidebar.text_input("Password", type="password")
        m = st.sidebar.text_input("Mobile (optional)")
        if st.sidebar.button("Register"):
            ok, msg = create_user(u, p, m)
            st.sidebar.success(msg if ok else msg)
    elif act == "Login":
        u = st.sidebar.text_input("Username")
        p = st.sidebar.text_input("Password", type="password")
        if st.sidebar.button("Login"):
            ok, res = authenticate_user(u, p)
            if ok:
                st.session_state.user_id = res
                st.session_state.username = u
                st.sidebar.success("Logged in!")
            else:
                st.sidebar.error(res)
else:
    st.sidebar.success(f"Logged in as {st.session_state.username}")
    if st.sidebar.button("Logout"):
        st.session_state.user_id = None
        st.session_state.username = None
        st.experimental_rerun()

use_model = st.sidebar.checkbox("Use full NLP models (optional)", value=False)
if use_model:
    load_models()

# -----------------------------
# MAIN INTERFACE
# -----------------------------
st.markdown("---")
page = st.radio("Navigation", ["Submit Comment", "Track Complaint"], horizontal=True)

SECTORS = [
    "Agriculture", "Mining and quarrying", "Manufacturing", "Electricity and gas",
    "Construction", "Trade", "Transport", "Business service"
]

if page == "Submit Comment":
    st.header("📝 Submit New Comment")
    sector = st.selectbox("Select Sector", SECTORS)
    comment = st.text_area("Enter your comment", height=150)
    if st.button("Submit Comment"):
        if not comment.strip():
            st.error("Please enter a valid comment.")
        else:
            uid = st.session_state.user_id if st.session_state.user_id else 0
            sentiment, score = predict_sentiment(comment, use_model)
            summary = summarize_text(comment, use_model)
            code = add_comment(uid, sector, comment, sentiment, summary)
            st.success(f"✅ Comment submitted successfully!")
            st.info(f"📌 Your Passcode: {code}")
            st.write(f"**Sentiment:** {sentiment.capitalize()} | **Summary:** {summary}")

elif page == "Track Complaint":
    st.header("🔎 Track Complaint")
    code = st.text_input("Enter your Passcode")
    if st.button("Track"):
        rec = get_comment_by_passcode(code.strip())
        if not rec:
            st.error("Passcode not found.")
        else:
            st.success("Record Found")
            st.write(f"**Sector:** {rec['sector']}")
            st.write(f"**Comment:** {rec['comment']}")
            st.write(f"**Summary:** {rec['summary']}")
            st.write(f"**Sentiment:** {rec['sentiment']}")
            st.write(f"**Status:** {rec['status']}")

# -----------------------------
# USER REPORT SECTION
# -----------------------------
if st.session_state.user_id:
    st.markdown("---")
    st.subheader("📄 Generate Summary Report")

    items = list_comments_for_user(st.session_state.user_id)
    if not items:
        st.info("No comments yet.")
    else:
        if st.button("Generate Word Report"):
            doc = Document()
            doc.add_heading("E-Consultation Summary Report", level=1)
            doc.add_paragraph(f"Generated for: {st.session_state.username}")
            doc.add_paragraph(f"Total Comments: {len(items)}")
            for i, it in enumerate(items, start=1):
                doc.add_heading(f"Comment #{i}", level=2)
                doc.add_paragraph(f"Sector: {it['sector']}")
                doc.add_paragraph(f"Passcode: {it['passcode']}")
                doc.add_paragraph(f"Comment: {it['comment']}")
                doc.add_paragraph(f"Sentiment: {it['sentiment']}")
                doc.add_paragraph(f"Summary: {it['summary']}")
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button(
                "⬇️ Download Word Report",
                data=buf,
                file_name="econsultation_summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

st.markdown("---")
st.caption("Built by Team Codezillas | Smart India Hackathon 2025")
